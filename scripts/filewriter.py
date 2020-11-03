#
#   Manege-Tool
#   Version 2
#   Author T.Wisotzki 2019
#

from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Pt


def Out(Trainingsday, out_dir):
    Columns = ['Datum', 'Aufwärmen', 'Lektion', 'Geräteturnen', 'Trampolin', 'Akrobatik']
    day = Trainingsday.name
    if day.__contains__("Montag"):
        Columns.extend(['Vertikaltuch', 'Slackline', 'Exra Thema'])
        day = "Montag"
        ort = "19:15 - 21:40 Zentrum"
    elif day.__contains__("Mittwoch"):
        Columns.extend(['Jonglieren', 'Slackline', 'Akro Bungee', 'Parkour'])
        day = "Mittwoch"
        ort = "19:15 - 21:40 Hönggerberg"
    elif day.__contains__("Donnerstag"):
        Columns.extend(['Jonglieren', 'Breakdance'])
        day = "Donnerstag"
        ort = "19:30 - 21:40 PHZ"
    else:
        print("Error: can not determine training-day")
        exit()
    Columns.append('Ersatz')

# formatting of the page
    document = Document()
    section = document.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.5)

# title and page setup
    document.add_heading('ASVZ Einsatzplan Manege ' + day, 0)
    p = document.add_paragraph('')
    p.add_run(ort).italic = True
    table = document.add_table(rows=1, cols=len(Columns))
    heading_cells = table.rows[0].cells
    for i in range(len(Columns)):
        heading_cells[i].text = Columns[i]

# setting the table contents
    for training in Trainingsday.trainings:
        row_cells = table.add_row().cells
        row_cells[0].text = training.date
        row_cells[3].text = training.GETU
        row_cells[4].text = training.TRA
        row_cells[5].text = training.AKRO
        if day == "Montag":
            row_cells[6].text = training.VERT
            row_cells[7].text = training.SLACK
            row_cells[9].text = training.CAN
        elif day == "Mittwoch":
            row_cells[6].text = training.JONG
            row_cells[7].text = training.SLACK
            row_cells[9].text = training.PARC
            row_cells[10].text = training.CAN
        elif day == "Donnerstag":
            row_cells[6].text = training.JONG
            row_cells[7].text = training.BREAK
            row_cells[8].text = training.CAN
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Arial'
                    font.size = Pt(8)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name = 'Arial'
                font.size = Pt(10)
                font.bold = True
    nam = "prov_Einteilung_" + day + ".docx"
    try:
        document.save(out_dir + nam)
    except Exception as e:
        print("  ERROR: specified output directory is not valid")
        exit()
